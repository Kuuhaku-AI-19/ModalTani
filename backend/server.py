from fastapi import FastAPI, APIRouter, UploadFile, File, Form, HTTPException
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import io
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
import math
import re
from datetime import datetime, timezone
from fastapi.responses import StreamingResponse
import json

try:
    from pypdf import PdfReader
    HAS_PDF = True
except Exception:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

# Qdrant + FastEmbed RAG engine (self-hosted, embedded)
try:
    from rag_engine import (
        semantic_chunk as rag_semantic_chunk,
        upsert_docs as rag_upsert,
        delete_doc as rag_delete,
        clear_collection as rag_clear,
        search as rag_search,
        hybrid_search as rag_hybrid_search,
        chunk_quality_report as rag_chunk_quality,
        is_ready as rag_is_ready,
        status as rag_status,
        count_points as rag_count,
    )
    HAS_RAG = True
except Exception as _rag_err:
    logging.warning(f"rag_engine not available: {_rag_err}")
    HAS_RAG = False


ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Try importing emergentintegrations for Gemini RAG
EMERGENT_LLM_KEY = os.environ.get("EMERGENT_LLM_KEY", "")
try:
    from emergentintegrations.llm.chat import LlmChat, UserMessage, TextDelta, StreamDone
    HAS_LLM_LIB = True
except Exception as e:
    logging.warning(f"Could not import emergentintegrations: {e}")
    HAS_LLM_LIB = False


# ----------------------------------------------------
# MODELS
# ----------------------------------------------------
class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    nama: str
    role: str # "petani" or "admin_bank"
    no_hp: str
    desa: Optional[str] = "Desa Sukamaju"
    kecamatan: Optional[str] = "Kec. Lembang"
    kabupaten: Optional[str] = "Kab. Bandung Barat"
    komoditas: Optional[str] = "Padi Sawah"
    avatar_url: Optional[str] = None

class QuizOption(BaseModel):
    id: str
    text: str

class QuizQuestion(BaseModel):
    id: str
    question: str
    options: List[QuizOption]
    correct_option_id: str
    explanation: str

class LearningModule(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    plek_category: str # "1. Pengelolaan Keuangan Rumah Tangga", "2. Pengelolaan Keuangan Usaha Tani", "3. Produk & Layanan Keuangan Formal"
    plek_category_number: int
    title: str
    subtitle: str
    duration_minutes: int
    media_type: str # "video" or "audio"
    media_url: str
    thumbnail_url: str
    summary_points: List[str]
    full_text_notes: str
    quiz: List[QuizQuestion]

class PretestQuestion(BaseModel):
    id: str
    question: str
    options: List[QuizOption]
    correct_option_id: str
    category: str

class PretestSubmit(BaseModel):
    user_id: str
    answers: Dict[str, str] # question_id -> option_id

class ModuleProgressUpdate(BaseModel):
    user_id: str
    module_id: str
    status: str # "completed", "in_progress"
    quiz_score: int # 0 to 100
    answers: Optional[Dict[str, str]] = None

class HarvestRecord(BaseModel):
    musim: str
    tahun: int
    volume_ton: float
    pendapatan_rp: int

class DocumentChecklist(BaseModel):
    ktp: bool = False
    kk: bool = False
    nib_atau_sku: bool = False
    sppt_pbb_atau_surat_lahan: bool = False
    buku_tabungan: bool = False
    foto_lahan: bool = False
    bpjs_ketenagakerjaan: bool = False  # required for KUR Kecil (> 100 jt)
    agunan_tambahan: bool = False  # SHM/BPKB required for KUR Kecil

class FarmerProfile(BaseModel):
    model_config = ConfigDict(extra="ignore")
    user_id: str
    nama: str
    nik: Optional[str] = "3201************"
    no_hp: str
    desa: str
    kecamatan: str
    kabupaten: str
    provinsi: str
    komoditas: str
    luas_lahan_ha: float
    status_lahan: str
    lama_bertani_tahun: int
    estimasi_pendapatan_musim_rp: int
    riwayat_panen: List[HarvestRecord]
    dokumen: DocumentChecklist
    pretest_completed: bool = False
    pretest_score: int = 0
    recommended_level: str = "Dasar"
    # Human supervisor judgment
    judgment_status: str = "pending"  # pending | approved | rejected
    judgment_note: Optional[str] = ""
    judged_at: Optional[str] = None
    judged_by: Optional[str] = None
    webinar_invited: bool = False
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class CreditScoreBreakdown(BaseModel):
    luas_lahan_poin: float
    riwayat_panen_poin: float
    edukasi_poin: float
    dokumen_poin: float
    total_score: float
    kategori: str
    badge_color: str
    rekomendasi_plafon: str
    rekomendasi_tindak_lanjut: List[str]
    auditable_factors: List[str]
    kur_tier: str = "KUR Mikro"  # "KUR Super Mikro" | "KUR Mikro" | "KUR Kecil"
    needs_one_on_one: bool = False  # True when KUR Kecil (> 100 jt)
    kur_kecil_gaps: List[str] = []  # missing BPJS TK / agunan for KUR Kecil

class JudgmentRequest(BaseModel):
    decision: str  # "approved" | "rejected" | "pending"
    note: Optional[str] = ""
    judged_by: Optional[str] = "Supervisor Admin"

class WebinarEvent(BaseModel):
    id: Optional[str] = None
    judul: str
    deskripsi: str
    tanggal: str  # ISO datetime
    durasi_menit: int = 90
    lokasi: str = "Online — Zoom Meeting"
    link_daftar: Optional[str] = ""
    pembicara: List[Dict[str, str]] = []  # [{nama, jabatan, institusi}]
    invited_user_ids: List[str] = []
    status: str = "upcoming"  # upcoming | completed | cancelled
    created_at: Optional[str] = None

class BankProduct(BaseModel):
    id: Optional[str] = None
    bank_nama: str
    bank_logo_url: Optional[str] = ""
    nama_produk: str
    tier: str = "KUR Mikro"  # KUR Super Mikro / KUR Mikro / KUR Kecil
    bunga_efektif_tahun: float = 6.0
    plafon_min_rp: int = 0
    plafon_max_rp: int = 100_000_000
    tenor_bulan_max: int = 36
    highlight: Optional[str] = ""  # "Yarnen tersedia", "Tanpa agunan sd 100 jt"
    link: Optional[str] = ""
    kontak_mantri: Optional[str] = ""
    active: bool = True
    created_at: Optional[str] = None

class KurDocument(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    topik: str
    judul: str
    kategori: str
    isi_teks: str
    sumber_nama: str
    sumber_link: str
    pasal_rujukan: Optional[str] = None

class ChatMessage(BaseModel):
    role: str # "user" or "assistant"
    text: str
    sumber_rujukan: Optional[List[str]] = None
    sumber_detail: Optional[List[Dict[str, str]]] = None
    timestamp: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ChatRequest(BaseModel):
    user_id: str
    message: str


# ----------------------------------------------------
# CREDIT SCORING RULE-BASED ENGINE (Auditable & Deterministic)
# ----------------------------------------------------
def calculate_farmer_credit_score(profile: dict, completed_modules_count: int, avg_quiz_score: float) -> CreditScoreBreakdown:
    # 1. Luas Lahan & Status Lahan (Max 25 Poin)
    luas = float(profile.get("luas_lahan_ha", 0))
    status_lahan = profile.get("status_lahan", "Milik Sendiri")
    
    if luas >= 2.0:
        lahan_pts = 22.0
    elif luas >= 1.0:
        lahan_pts = 18.0
    elif luas >= 0.5:
        lahan_pts = 14.0
    else:
        lahan_pts = 10.0
    
    # Status kepemilikan
    if status_lahan == "Milik Sendiri":
        lahan_pts += 3.0
    elif status_lahan == "Sewa":
        lahan_pts += 2.0
    else:
        lahan_pts += 1.0
    lahan_pts = min(25.0, lahan_pts)

    # 2. Riwayat Panen & Lama Usaha (Max 25 Poin)
    lama_bertani = int(profile.get("lama_bertani_tahun", 1))
    pendapatan = int(profile.get("estimasi_pendapatan_musim_rp", 0))
    
    panen_pts = 0.0
    if lama_bertani >= 5:
        panen_pts += 12.0
    elif lama_bertani >= 2:
        panen_pts += 8.0
    else:
        panen_pts += 4.0
        
    if pendapatan >= 25000000:
        panen_pts += 13.0
    elif pendapatan >= 15000000:
        panen_pts += 10.0
    elif pendapatan >= 8000000:
        panen_pts += 7.0
    else:
        panen_pts += 4.0
    panen_pts = min(25.0, panen_pts)

    # 3. Status Penyelesaian Edukasi ModalTani (Max 25 Poin)
    # 3 modules in total
    if completed_modules_count >= 3:
        edu_pts = 22.0
    elif completed_modules_count == 2:
        edu_pts = 15.0
    elif completed_modules_count == 1:
        edu_pts = 8.0
    else:
        edu_pts = 0.0
        
    if avg_quiz_score >= 80 and completed_modules_count > 0:
        edu_pts += 3.0
    elif avg_quiz_score >= 60 and completed_modules_count > 0:
        edu_pts += 1.5
    edu_pts = min(25.0, edu_pts)

    # 4. Kelengkapan Dokumen Administratif (Max 25 Poin)
    docs = profile.get("dokumen", {})
    doc_pts = 0.0
    if docs.get("ktp", False):
        doc_pts += 5.0
    if docs.get("kk", False):
        doc_pts += 5.0
    if docs.get("nib_atau_sku", False):
        doc_pts += 10.0  # Crucial legal basis for KUR
    if docs.get("sppt_pbb_atau_surat_lahan", False):
        doc_pts += 3.0
    if docs.get("buku_tabungan", False):
        doc_pts += 2.0
    doc_pts = min(25.0, doc_pts)

    total = round(lahan_pts + panen_pts + edu_pts + doc_pts, 1)
    
    # Categorization & Actionable Recommendations
    rekomendasi = []
    auditable = []
    
    auditable.append(f"Luas lahan {luas} Ha ({status_lahan}): {lahan_pts}/25 pts")
    auditable.append(f"Pengalaman {lama_bertani} thn & omzet Rp {pendapatan:,.0f}/musim: {panen_pts}/25 pts")
    auditable.append(f"Literasi ModalTani ({completed_modules_count}/3 modul, kuis avg {avg_quiz_score:.0f}%): {edu_pts}/25 pts")
    auditable.append(f"Kelengkapan berkas legal (KTP/KK/NIB/SPPT): {doc_pts}/25 pts")

    if total >= 70.0:
        kategori = "Layak Direkomendasikan"
        badge_color = "green"
        rekomendasi_plafon = "KUR Mikro Rp 25.000.000 - Rp 100.000.000 (Bunga 6% p.a.)"
        rekomendasi.append("Siap diajukan ke Bank Himbara untuk verifikasi berkas & survei lapangan")
        if not docs.get("nib_atau_sku"):
            rekomendasi.append("Segera terbitkan NIB online via OSS atau minta Surat Keterangan Usaha (SKU) Desa")
    elif total >= 40.0:
        kategori = "Perlu Pendampingan Lanjutan"
        badge_color = "amber"
        rekomendasi_plafon = "KUR Super Mikro s.d Rp 10.000.000 (Bunga 3% p.a.) / Pendampingan KUR Mikro"
        if completed_modules_count < 3:
            rekomendasi.append(f"Selesaikan {3 - completed_modules_count} modul edukasi tersisa untuk menaikkan skor hingga +{25 - edu_pts:.0f} poin")
        if not docs.get("nib_atau_sku"):
            rekomendasi.append("Lengkapi legalitas usaha tani (SKU dari Kantor Desa / NIB) (+10 poin)")
        if not docs.get("ktp") or not docs.get("kk"):
            rekomendasi.append("Lengkapi fotokopi KTP & Kartu Keluarga (+10 poin)")
    else:
        kategori = "Belum Layak — Edukasi Dulu"
        badge_color = "red"
        rekomendasi_plafon = "Belum disarankan pengajuan kredit formal (Fokus peningkatan kapasitas)"
        rekomendasi.append("Wajib selesaikan Modul 1 & 2 Pengelolaan Keuangan Dasar ModalTani")
        rekomendasi.append("Dampingi pembuatan pembukuan usaha tani sederhana bersama Penyuluh (PPL)")
        rekomendasi.append("Urus kelengkapan data kependudukan dan surat keterangan garapan dari desa")

    # KUR tier detection: KUR Kecil requires higher qualification + BPJS TK + agunan tambahan
    # Triggers KUR Kecil when: skor >= 80 AND luas >= 2 Ha AND omzet >= 20jt/musim
    kur_tier = "KUR Mikro"
    needs_one_on_one = False
    kur_kecil_gaps: List[str] = []
    if total >= 80.0 and luas >= 2.0 and pendapatan >= 20_000_000:
        kur_tier = "KUR Kecil"
        needs_one_on_one = True
        rekomendasi_plafon = "KUR Kecil Rp 101.000.000 - Rp 500.000.000 (Bunga 6% p.a., wajib agunan tambahan)"
        rekomendasi.insert(0, "🎯 Kandidat KUR Kecil (> Rp 100 Juta) — layak konsultasi 1-on-1 dengan Kepala Unit/Cabang Bank")
        if not docs.get("bpjs_ketenagakerjaan"):
            kur_kecil_gaps.append("BPJS Ketenagakerjaan aktif (wajib untuk KUR Kecil)")
            rekomendasi.append("Daftarkan BPJS Ketenagakerjaan (BPJS TK) — syarat wajib KUR Kecil > Rp 100 Juta")
        if not docs.get("agunan_tambahan"):
            kur_kecil_gaps.append("Agunan tambahan (SHM / BPKB kendaraan)")
            rekomendasi.append("Siapkan agunan tambahan berupa Sertifikat Hak Milik (SHM) atau BPKB kendaraan")
    elif total >= 40.0 and total < 70.0:
        kur_tier = "KUR Super Mikro"
    elif total < 40.0:
        kur_tier = "Belum Direkomendasikan"

    auditable.append(f"Tier akhir: {kur_tier}" + (" • butuh konsultasi 1-on-1" if needs_one_on_one else ""))

    return CreditScoreBreakdown(
        luas_lahan_poin=lahan_pts,
        riwayat_panen_poin=panen_pts,
        edukasi_poin=edu_pts,
        dokumen_poin=doc_pts,
        total_score=total,
        kategori=kategori,
        badge_color=badge_color,
        rekomendasi_plafon=rekomendasi_plafon,
        rekomendasi_tindak_lanjut=rekomendasi,
        auditable_factors=auditable,
        kur_tier=kur_tier,
        needs_one_on_one=needs_one_on_one,
        kur_kecil_gaps=kur_kecil_gaps,
    )


# ----------------------------------------------------
# RAG RETRIEVAL ENGINE
# ----------------------------------------------------
async def retrieve_relevant_kur_docs(query: str, top_k: int = 3) -> List[dict]:
    # 1) Hybrid search (BM25 + dense) preferred
    if HAS_RAG and rag_is_ready():
        try:
            corpus = await db.kur_documents.find({}, {"_id": 0}).to_list(500)
            hits = rag_hybrid_search(query, corpus, top_k=top_k)
            if hits:
                cleaned = []
                for h in hits:
                    cleaned.append({k: v for k, v in h.items() if not k.startswith("_")})
                return cleaned
        except Exception as e:
            logging.warning(f"Hybrid search failed, falling back to keyword: {e}")

    # 2) Fallback: keyword-weighted retrieval over mongo
    docs = await db.kur_documents.find({}, {"_id": 0}).to_list(200)
    if not docs:
        return []
    query_tokens = set(re.findall(r'\w+', query.lower()))
    keywords_weight = {
        "syarat": 3, "dokumen": 3, "ktp": 3, "sku": 3, "nib": 3,
        "bunga": 3, "persen": 2, "6%": 4, "3%": 4, "agunan": 3, "jaminan": 3,
        "plafon": 3, "pinjaman": 2, "mikro": 2, "sewa": 3, "garapan": 3,
        "tahapan": 2, "alur": 2, "bri": 2, "mandiri": 2, "bni": 2, "bsi": 2,
        "yarnen": 4, "panen": 2, "angsuran": 2, "survei": 2, "mantri": 2
    }
    scored = []
    for doc in docs:
        text = f"{doc.get('topik', '')} {doc.get('judul', '')} {doc.get('isi_teks', '')}".lower()
        score = 0.0
        for token in query_tokens:
            if token in text:
                weight = keywords_weight.get(token, 1.0)
                score += weight * (text.count(token) + 1)
        scored.append((score, doc))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [d for _, d in scored[:top_k]]


async def generate_rag_answer(query: str, relevant_docs: List[dict]) -> tuple[str, List[str], List[Dict[str, str]]]:
    context_blocks = []
    sumber_names = []
    sumber_details = []
    
    for doc in relevant_docs:
        context_blocks.append(
            f"### Dokumen: {doc.get('judul')}\n"
            f"Sumber: {doc.get('sumber_nama')} ({doc.get('pasal_rujukan', 'Ketentuan Resmi')})\n"
            f"Isi Ketentuan:\n{doc.get('isi_teks')}\n"
        )
        sumber_names.append(doc.get('sumber_nama', 'Sumber Resmi'))
        sumber_details.append({
            "nama": doc.get('sumber_nama', 'Dokumen Resmi'),
            "judul": doc.get('judul', 'Pedoman KUR Pertanian'),
            "link": doc.get('sumber_link', '#'),
            "pasal": doc.get('pasal_rujukan', 'Standar Operasional Prosedur')
        })
        
    context_str = "\n".join(context_blocks)
    
    system_prompt = (
        "Anda adalah 'Asisten KUR ModalTani', pakar perbankan dan KUR pertanian yang ramah, santun, dan berbicara dalam bahasa Indonesia yang sangat mudah dipahami oleh petani desa ('bahasa tani' yang bersahabat dan lugas).\n"
        "PANDUAN UTAMA:\n"
        "1. Jawab HANYA berdasarkan fakta dari Dokumen Resmi yang disediakan di bagian Konteks.\n"
        "2. Jangan pernah berhalusinasi atau mengarang syarat yang tidak tertulis.\n"
        "3. Gunakan sapaan hangat seperti 'Bapak/Ibu Petani'.\n"
        "4. Jelaskan angka dan prosedur secara praktis (misal: bunga 6% per tahun, skema Yarnen/bayar saat panen, agunan pokok tanaman/objek usaha).\n"
        "5. Di akhir jawaban, berikan bagian 'Rekomendasi Langkah Praktis' singkat (1-3 butir).\n"
        "6. Selalu sebutkan bahwa jawaban ini merujuk pada regulasi OJK & Permenko Perekonomian yang berlaku."
    )
    
    user_prompt = f"Pertanyaan Petani: {query}\n\n--- KONTEKS DOKUMEN RESMI ---\n{context_str}\n\nBerikan jawaban ramah dan terstruktur:"
    
    # Try with Gemini LLM
    if HAS_LLM_LIB and EMERGENT_LLM_KEY:
        try:
            chat = LlmChat(
                api_key=EMERGENT_LLM_KEY,
                session_id=str(uuid.uuid4()),
                system_message=system_prompt
            ).with_model("gemini", "gemini-3-flash-preview")
            
            response = await chat.send_message(UserMessage(text=user_prompt))
            # response may be either a plain string or an object with .text
            response_text = None
            if isinstance(response, str):
                response_text = response
            elif hasattr(response, 'text'):
                response_text = response.text
            if response_text:
                return response_text, list(set(sumber_names)), sumber_details
        except Exception as e:
            logging.warning(f"Gemini API call failed, using high-fidelity fallback: {e}")
            
    # Intelligent high-fidelity fallback generator if offline
    fallback_text = build_fallback_agricultural_response(query, relevant_docs)
    return fallback_text, list(set(sumber_names)), sumber_details


def build_fallback_agricultural_response(query: str, docs: List[dict]) -> str:
    q = query.lower()
    doc_excerpts = "\n".join([f"• **{d.get('judul')}** ({d.get('sumber_nama')}): {d.get('isi_teks')[:180]}..." for d in docs])
    
    if "syarat" in q or "dokumen" in q or "berkas" in q or "butuh apa" in q:
        return (
            "Halo Bapak/Ibu Petani! Berdasarkan pedoman resmi KUR Sektor Pertanian, berikut syarat utama yang perlu disiapkan:\n\n"
            "1. **Identitas Diri**: Fotokopi KTP Elektronik & Kartu Keluarga (KK).\n"
            "2. **Legalitas Usaha Tani**: Surat Keterangan Usaha (SKU) dari Kantor Desa/Kelurahan atau Nomor Induk Berusaha (NIB) yang bisa dibuat secara gratis.\n"
            "3. **Pengalaman Bertani**: Usaha tani sudah aktif berjalan minimal 6 bulan.\n"
            "4. **Status Pinjaman**: Tidak sedang memiliki kredit komersial produktif lain (kecuali kredit konsumtif seperti KPR/Kredit Motor yang lancar).\n\n"
            "📋 **Langkah Praktis Berikutnya:**\n"
            "1. Datang ke Balai Desa untuk meminta surat pengantar SKU Usaha Tani.\n"
            "2. Selesaikan modul belajar di ModalTani untuk memperoleh sertifikat literasi sebagai nilai tambah penilaian bank.\n"
            "3. Bawa berkas ke unit bank penyalur (BRI/Mandiri/BNI/BSI) terdekat."
        )
    elif "bunga" in q or "persen" in q or "biaya" in q or "yarnen" in q or "bayar" in q:
        return (
            "Halo Bapak/Ibu! Untuk Kredit Usaha Rakyat (KUR) Pertanian:\n\n"
            "1. **Suku Bunga Subsidi**: Sangat ringan yaitu **6% efektif per tahun** (sekitar 0,2% per bulan) karena disubsidi oleh pemerintah.\n"
            "2. **Skema Pembayaran Yarnen (Bayar Setelah Panen)**: Petani tanaman pangan (seperti padi dan jagung) bisa memilih skema angsuran sekaligus saat masa panen tiba, sehingga tidak terbebani cicilan bulanan saat masa tanam.\n"
            "3. **Bebas Biaya Administrasi & Provisi**: Untuk KUR Mikro tidak dikenakan potongan biaya tersembunyi.\n\n"
            "🌾 **Saran ModalTani:** Manfaatkan kalkulator arus kas di Modul 2 untuk menghitung estimasi pendapatan panen sebelum menentukan nominal pinjaman."
        )
    elif "agunan" in q or "jaminan" in q or "sertifikat" in q or "gadai" in q:
        return (
            "Kabar baik untuk Bapak/Ibu Petani! Berdasarkan Peraturan Menteri Koordinator Perekonomian RI:\n\n"
            "1. **Plafon s.d Rp 100 Juta (KUR Mikro)**: **TIDAK DIWAJIBKAN agunan tambahan** (seperti sertifikat tanah/BPKB). Agunan pokoknya adalah usaha tani dan komoditas tanaman itu sendiri.\n"
            "2. **Plafon di atas Rp 100 Juta (KUR Kecil)**: Memerlukan agunan tambahan sesuai ketentuan penilaian bank penyalur.\n\n"
            "💡 **Penting:** Pastikan pembukuan panen dan luas garapan dicatat dengan rapi agar pihak bank mantri yakin dengan kapasitas produksi Bapak/Ibu."
        )
    elif "sewa" in q or "garap" in q or "adat" in q or "bukan tanah sendiri" in q:
        return (
            "Halo Bapak/Ibu Petani! Petani penggarap atau penyewa lahan **tetap berhak dan bisa mengajukan KUR**.\n\n"
            "Ketentuannya adalah:\n"
            "1. Menyertakan **Surat Perjanjian Sewa Lahan** atau **Surat Keterangan Garap** yang diketahui dan ditandatangani oleh Kepala Desa/Lurah.\n"
            "2. Masa sewa/garap sebaiknya mencakup minimal satu siklus masa panen komoditas yang diajukan.\n\n"
            "🤝 **Langkah:** Mintalah format surat garap/sewa sederhana ke perangkat desa setempat sebelum datang ke bank."
        )
    else:
        return (
            f"Halo Bapak/Ibu Petani! Terima kasih atas pertanyaannya mengenai: *{query}*.\n\n"
            f"Berdasarkan himpunan dokumen resmi KUR Pertanian:\n"
            f"{doc_excerpts}\n\n"
            "🌟 **Saran ModalTani:** Program KUR dirancang agar modal kerja petani tercukupi dengan bunga terjangkau 6%. Jika ada hal teknis lain yang ingin ditanyakan seputar berkas atau skema bank, silakan tanyakan langsung ke asisten ini!"
        )


# ----------------------------------------------------
# ENDPOINTS
# ----------------------------------------------------
@api_router.get("/")
async def root():
    return {
        "status": "online",
        "app": "ModalTani API",
        "description": "Edutech & AI KUR Advisory Platform for Indonesian Farmers",
        "version": "1.0.0"
    }

# 1. Users & Demo Auth
@api_router.get("/users")
async def get_users():
    users = await db.users.find({}, {"_id": 0}).to_list(100)
    return users

@api_router.post("/auth/demo-login")
async def demo_login(payload: dict):
    role = payload.get("role", "petani")
    user_id = payload.get("user_id")
    nama = payload.get("nama", "Pengguna Demo")
    
    if user_id:
        user = await db.users.find_one({"id": user_id}, {"_id": 0})
        if user:
            return {"success": True, "user": user}
            
    # Find first user of this role
    user = await db.users.find_one({"role": role}, {"_id": 0})
    if user:
        return {"success": True, "user": user}
        
    # Create fallback demo user
    new_user = {
        "id": str(uuid.uuid4()),
        "nama": nama,
        "role": role,
        "no_hp": payload.get("no_hp", "081234567890"),
        "desa": "Desa Sukamaju",
        "kecamatan": "Kec. Lembang",
        "kabupaten": "Kab. Bandung Barat",
        "komoditas": "Padi Sawah"
    }
    await db.users.insert_one(new_user.copy())
    return {"success": True, "user": new_user}


# 2. Learning Modules (PLEK Kementan)
@api_router.get("/modules")
async def get_modules():
    modules = await db.learning_modules.find({}, {"_id": 0}).sort("plek_category_number", 1).to_list(100)
    return modules

@api_router.get("/modules/{module_id}")
async def get_module_detail(module_id: str):
    mod = await db.learning_modules.find_one({"id": module_id}, {"_id": 0})
    if not mod:
        return {"error": "Modul tidak ditemukan"}
    return mod


# 3. Pretest (Literacy Assessment)
@api_router.get("/pretest")
async def get_pretest():
    questions = await db.pretest_questions.find({}, {"_id": 0}).to_list(20)
    return questions

@api_router.post("/pretest/submit")
async def submit_pretest(data: PretestSubmit):
    questions = await db.pretest_questions.find({}, {"_id": 0}).to_list(20)
    q_dict = {q["id"]: q for q in questions}
    
    correct_count = 0
    results = []
    for q_id, chosen_option in data.answers.items():
        q_data = q_dict.get(q_id)
        if q_data:
            is_correct = (chosen_option == q_data.get("correct_option_id"))
            if is_correct:
                correct_count += 1
            results.append({
                "question_id": q_id,
                "is_correct": is_correct,
                "chosen": chosen_option,
                "correct": q_data.get("correct_option_id")
            })
            
    total_q = max(1, len(questions))
    score_pct = int((correct_count / total_q) * 100)
    recommended_level = "Menengah (Fokus Produk KUR)" if score_pct >= 60 else "Dasar (Kelola Kas & HPP Dulu)"
    
    # Update profile
    await db.farmer_profiles.update_one(
        {"user_id": data.user_id},
        {"$set": {
            "pretest_completed": True,
            "pretest_score": score_pct,
            "recommended_level": recommended_level
        }},
        upsert=True
    )
    
    return {
        "score": score_pct,
        "correct_count": correct_count,
        "total_questions": total_q,
        "recommended_level": recommended_level,
        "results": results
    }


# 4. Learning Progress
@api_router.get("/learning-progress/{user_id}")
async def get_learning_progress(user_id: str):
    records = await db.learning_progress.find({"user_id": user_id}, {"_id": 0}).to_list(100)
    modules = await db.learning_modules.find({}, {"_id": 0}).to_list(100)
    
    total_modules = len(modules)
    completed_count = sum(1 for r in records if r.get("status") == "completed")
    quiz_scores = [r.get("quiz_score", 0) for r in records if r.get("quiz_score") is not None]
    avg_quiz_score = sum(quiz_scores) / len(quiz_scores) if quiz_scores else 0
    
    progress_pct = int((completed_count / total_modules * 100)) if total_modules > 0 else 0
    
    return {
        "user_id": user_id,
        "total_modules": total_modules,
        "completed_count": completed_count,
        "progress_percentage": progress_pct,
        "avg_quiz_score": avg_quiz_score,
        "records": records
    }

@api_router.post("/learning-progress/update")
async def update_learning_progress(data: ModuleProgressUpdate):
    record = {
        "user_id": data.user_id,
        "module_id": data.module_id,
        "status": data.status,
        "quiz_score": data.quiz_score,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    await db.learning_progress.update_one(
        {"user_id": data.user_id, "module_id": data.module_id},
        {"$set": record},
        upsert=True
    )
    
    return {"success": True, "record": record}


# 5. KUR Official Documents Knowledge Base
@api_router.get("/kur-docs")
async def get_kur_docs():
    docs = await db.kur_documents.find({}, {"_id": 0}).to_list(100)
    return docs

class KurDocCreate(BaseModel):
    topik: str
    judul: str
    kategori: str = "Regulasi Umum"
    isi_teks: str
    sumber_nama: str
    sumber_link: Optional[str] = "#"
    pasal_rujukan: Optional[str] = ""

@api_router.post("/kur-docs")
async def create_kur_doc(doc: KurDocCreate):
    new_doc = {
        "id": f"doc-{uuid.uuid4().hex[:8]}",
        "topik": doc.topik,
        "judul": doc.judul,
        "kategori": doc.kategori,
        "isi_teks": doc.isi_teks,
        "sumber_nama": doc.sumber_nama,
        "sumber_link": doc.sumber_link or "#",
        "pasal_rujukan": doc.pasal_rujukan or "Ketentuan Resmi",
    }
    await db.kur_documents.insert_one(new_doc.copy())
    if HAS_RAG:
        rag_upsert([new_doc])
    return {"success": True, "doc": new_doc}

@api_router.delete("/kur-docs/{doc_id}")
async def delete_kur_doc(doc_id: str):
    result = await db.kur_documents.delete_one({"id": doc_id})
    if HAS_RAG:
        rag_delete(doc_id)
    return {"success": result.deleted_count > 0}


@api_router.get("/rag/status")
async def get_rag_status():
    if not HAS_RAG:
        return {"available": False, "reason": "rag_engine not installed"}
    return {"available": True, **rag_status(), "points_in_qdrant": rag_count()}


@api_router.post("/rag/reindex")
async def reindex_rag():
    if not HAS_RAG:
        return {"success": False, "reason": "rag_engine not installed"}
    docs = await db.kur_documents.find({}, {"_id": 0}).to_list(1000)
    rag_clear()
    n = rag_upsert(docs)
    return {"success": True, "docs_indexed": n, "mongo_source_count": len(docs)}


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    if not HAS_PDF:
        raise HTTPException(status_code=500, detail="Library pypdf tidak terpasang di server")
    reader = PdfReader(io.BytesIO(file_bytes))
    parts = []
    for page in reader.pages:
        try:
            txt = page.extract_text() or ""
            if txt.strip():
                parts.append(txt.strip())
        except Exception:
            continue
    return "\n\n".join(parts)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    if not HAS_DOCX:
        raise HTTPException(status_code=500, detail="Library python-docx tidak terpasang di server")
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # also include table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n\n".join(parts)


def _chunk_text(text: str, target_size: int = 800, overlap: int = 120) -> List[str]:
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    if not text:
        return []
    if len(text) <= target_size:
        return [text]

    # Split by sentences first, then pack into ~target_size chunks
    sentences = re.split(r'(?<=[\.\?\!])\s+', text)
    chunks = []
    current = ""
    for sent in sentences:
        if len(current) + len(sent) + 1 <= target_size:
            current = (current + " " + sent).strip() if current else sent
        else:
            if current:
                chunks.append(current)
            # Add overlap tail of previous chunk into new one for RAG continuity
            tail = current[-overlap:] if current and overlap > 0 else ""
            current = (tail + " " + sent).strip() if tail else sent
            # If single sentence is longer than target_size, hard-split it
            while len(current) > target_size * 1.5:
                chunks.append(current[:target_size])
                current = current[target_size - overlap:]
    if current:
        chunks.append(current)
    return chunks


@api_router.post("/kur-docs/upload")
async def upload_kur_doc_file(
    file: UploadFile = File(...),
    kategori: str = Form("Regulasi Umum"),
    sumber_nama: str = Form(""),
    sumber_link: str = Form("#"),
    pasal_rujukan: str = Form(""),
    target_chunk_size: int = Form(800),
):
    filename = file.filename or "uploaded.pdf"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in {"pdf", "docx"}:
        raise HTTPException(status_code=400, detail="Format tidak didukung. Gunakan PDF atau DOCX.")

    contents = await file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="File kosong")

    if ext == "pdf":
        raw_text = _extract_text_from_pdf(contents)
    else:
        raw_text = _extract_text_from_docx(contents)

    if not raw_text.strip():
        raise HTTPException(status_code=400, detail="Tidak ada teks yang bisa diekstrak dari file ini")

    # Prefer semantic chunking via Qdrant/SemanticChunker; fallback to naive
    if HAS_RAG:
        chunks = rag_semantic_chunk(raw_text)
    else:
        chunks = _chunk_text(raw_text, target_size=max(300, min(2000, target_chunk_size)))
    if not chunks:
        raise HTTPException(status_code=400, detail="Ekstraksi menghasilkan 0 chunk")

    base_title = filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").strip().title()
    sumber = sumber_nama or f"Dokumen Unggahan: {filename}"

    inserted = []
    for idx, chunk in enumerate(chunks, start=1):
        entry = {
            "id": f"doc-{uuid.uuid4().hex[:8]}",
            "topik": f"{base_title} — Bagian {idx}/{len(chunks)}",
            "judul": f"{base_title} (Bagian {idx})",
            "kategori": kategori or "Regulasi Umum",
            "isi_teks": chunk,
            "sumber_nama": sumber,
            "sumber_link": sumber_link or "#",
            "pasal_rujukan": pasal_rujukan or f"Halaman {idx}",
            "uploaded_at": datetime.now(timezone.utc).isoformat(),
            "source_file": filename,
        }
        inserted.append(entry)

    if inserted:
        await db.kur_documents.insert_many([e.copy() for e in inserted])
    indexed = 0
    if HAS_RAG and inserted:
        indexed = rag_upsert(inserted)

    # Chunk quality (cosine similarity between adjacent chunks)
    quality_report = []
    if HAS_RAG and len(chunks) >= 2:
        try:
            quality_report = rag_chunk_quality(chunks)
        except Exception:
            pass

    return {
        "success": True,
        "filename": filename,
        "file_type": ext,
        "chunks_created": len(inserted),
        "vectors_indexed": indexed if HAS_RAG else 0,
        "used_semantic_chunker": HAS_RAG,
        "total_chars": sum(len(e["isi_teks"]) for e in inserted),
        "doc_ids": [e["id"] for e in inserted],
        "chunk_quality": quality_report,
        "preview_first_chunk": inserted[0]["isi_teks"][:250] + ("..." if len(inserted[0]["isi_teks"]) > 250 else "") if inserted else "",
    }


# 6. RAG AI KUR Advisory Chat
@api_router.post("/chat/ask")
async def chat_ask(req: ChatRequest):
    user_id = req.user_id
    user_query = req.message.strip()
    
    if not user_query:
        return {"error": "Pesan tidak boleh kosong"}
        
    # 1. Retrieve top matching official documents
    relevant_docs = await retrieve_relevant_kur_docs(user_query, top_k=3)
    
    # 2. Generate answer with Gemini RAG pipeline
    bot_answer, sources, source_details = await generate_rag_answer(user_query, relevant_docs)
    
    # 3. Store in DB chat history
    user_msg = {
        "role": "user",
        "text": user_query,
        "timestamp": datetime.now(timezone.utc).isoformat()
    }
    bot_msg = {
        "role": "assistant",
        "text": bot_answer,
        "sumber_rujukan": sources,
        "sumber_detail": source_details,
        "timestamp": datetime.now(timezone.utc).isoformat()
    }
    
    await db.chat_sessions.update_one(
        {"user_id": user_id},
        {"$push": {"messages": {"$each": [user_msg, bot_msg]}}},
        upsert=True
    )
    
    return {
        "answer": bot_answer,
        "sources": sources,
        "source_details": source_details,
        "disclaimer": "Dijawab otomatis berbasis dokumen resmi OJK & Permenko Perekonomian"
    }

@api_router.get("/chat/history/{user_id}")
async def get_chat_history(user_id: str):
    session = await db.chat_sessions.find_one({"user_id": user_id}, {"_id": 0})
    if not session or "messages" not in session:
        return {"user_id": user_id, "messages": []}
    return session

@api_router.delete("/chat/history/{user_id}")
async def clear_chat_history(user_id: str):
    await db.chat_sessions.delete_one({"user_id": user_id})
    return {"success": True, "message": "Riwayat chat berhasil dibersihkan"}


# 7. Bank Mitra CRM & Credit Scoring Dashboard
@api_router.get("/crm/farmers")
async def get_crm_farmers(
    desa: Optional[str] = None,
    komoditas: Optional[str] = None,
    kategori_skor: Optional[str] = None,
    search: Optional[str] = None
):
    profiles = await db.farmer_profiles.find({}, {"_id": 0}).to_list(200)
    
    # Enrich profiles with real-time calculated credit scores and progress
    enriched_farmers = []
    for prof in profiles:
        u_id = prof.get("user_id")
        # Get progress
        progress_records = await db.learning_progress.find({"user_id": u_id}, {"_id": 0}).to_list(100)
        completed_count = sum(1 for r in progress_records if r.get("status") == "completed")
        quiz_scores = [r.get("quiz_score", 0) for r in progress_records if r.get("quiz_score") is not None]
        avg_score = sum(quiz_scores) / len(quiz_scores) if quiz_scores else 0
        
        score_calc = calculate_farmer_credit_score(prof, completed_count, avg_score)
        
        # Filter criteria
        if desa and desa.lower() not in prof.get("desa", "").lower():
            continue
        if komoditas and komoditas.lower() not in prof.get("komoditas", "").lower():
            continue
        if kategori_skor and kategori_skor != "Semua" and score_calc.kategori != kategori_skor:
            continue
        if search:
            s = search.lower()
            match = (
                s in prof.get("nama", "").lower() or
                s in prof.get("desa", "").lower() or
                s in prof.get("komoditas", "").lower() or
                s in prof.get("no_hp", "").lower()
            )
            if not match:
                continue
                
        enriched_farmers.append({
            **prof,
            "modul_selesai": completed_count,
            "total_modul": 3,
            "credit_score": score_calc.model_dump()
        })
        
    # Sort by total score descending
    enriched_farmers.sort(key=lambda x: x["credit_score"]["total_score"], reverse=True)
    return enriched_farmers

@api_router.get("/crm/farmers/{user_id}")
async def get_crm_farmer_detail(user_id: str):
    prof = await db.farmer_profiles.find_one({"user_id": user_id}, {"_id": 0})
    if not prof:
        return {"error": "Profil petani tidak ditemukan"}
        
    progress_records = await db.learning_progress.find({"user_id": user_id}, {"_id": 0}).to_list(100)
    completed_count = sum(1 for r in progress_records if r.get("status") == "completed")
    quiz_scores = [r.get("quiz_score", 0) for r in progress_records if r.get("quiz_score") is not None]
    avg_score = sum(quiz_scores) / len(quiz_scores) if quiz_scores else 0
    
    score_calc = calculate_farmer_credit_score(prof, completed_count, avg_score)
    
    return {
        **prof,
        "modul_selesai": completed_count,
        "total_modul": 3,
        "credit_score": score_calc.model_dump(),
        "progress_detail": progress_records
    }

@api_router.put("/crm/farmers/{user_id}/documents")
async def update_farmer_documents(user_id: str, docs: DocumentChecklist):
    await db.farmer_profiles.update_one(
        {"user_id": user_id},
        {"$set": {"dokumen": docs.model_dump()}}
    )
    return {"success": True, "message": "Status verifikasi dokumen berhasil diperbarui"}

@api_router.post("/crm/farmers")
async def create_farmer_profile(profile: FarmerProfile):
    prof_dict = profile.model_dump()
    
    # Create associated user if not exist
    user_doc = {
        "id": profile.user_id,
        "nama": profile.nama,
        "role": "petani",
        "no_hp": profile.no_hp,
        "desa": profile.desa,
        "kecamatan": profile.kecamatan,
        "kabupaten": profile.kabupaten,
        "komoditas": profile.komoditas
    }
    await db.users.update_one({"id": profile.user_id}, {"$set": user_doc}, upsert=True)
    await db.farmer_profiles.update_one({"user_id": profile.user_id}, {"$set": prof_dict}, upsert=True)
    
    return {"success": True, "farmer": prof_dict}

@api_router.get("/crm/analytics")
async def get_crm_analytics():
    profiles = await db.farmer_profiles.find({}, {"_id": 0}).to_list(200)
    
    total_farmers = len(profiles)
    category_counts = {
        "Layak Direkomendasikan": 0,
        "Perlu Pendampingan Lanjutan": 0,
        "Belum Layak — Edukasi Dulu": 0
    }
    
    total_score_sum = 0.0
    total_lahan = 0.0
    total_potential_plafon = 0
    
    for prof in profiles:
        u_id = prof.get("user_id")
        progress_records = await db.learning_progress.find({"user_id": u_id}, {"_id": 0}).to_list(100)
        completed_count = sum(1 for r in progress_records if r.get("status") == "completed")
        quiz_scores = [r.get("quiz_score", 0) for r in progress_records if r.get("quiz_score") is not None]
        avg_score = sum(quiz_scores) / len(quiz_scores) if quiz_scores else 0
        
        score_calc = calculate_farmer_credit_score(prof, completed_count, avg_score)
        cat = score_calc.kategori
        if cat in category_counts:
            category_counts[cat] += 1
        total_score_sum += score_calc.total_score
        total_lahan += float(prof.get("luas_lahan_ha", 0))
        
        if cat == "Layak Direkomendasikan":
            total_potential_plafon += 50000000 # avg 50jt
        elif cat == "Perlu Pendampingan Lanjutan":
            total_potential_plafon += 10000000 # 10jt super mikro
            
    avg_score = round(total_score_sum / total_farmers, 1) if total_farmers > 0 else 0
    
    distribution_chart = [
        {"name": "Layak Direkomendasikan (>=70)", "value": category_counts["Layak Direkomendasikan"], "color": "#166534"},
        {"name": "Perlu Pendampingan (40-69)", "value": category_counts["Perlu Pendampingan Lanjutan"], "color": "#D97706"},
        {"name": "Belum Layak (<40)", "value": category_counts["Belum Layak — Edukasi Dulu"], "color": "#DC2626"}
    ]
    
    return {
        "total_farmers": total_farmers,
        "average_credit_score": avg_score,
        "total_lahan_ha": round(total_lahan, 2),
        "total_potential_financing_rp": total_potential_plafon,
        "category_counts": category_counts,
        "distribution_chart": distribution_chart
    }


# ----------------------------------------------------
# SUPERVISOR JUDGMENT ENDPOINTS
# ----------------------------------------------------
@api_router.patch("/crm/farmers/{user_id}/judgment")
async def submit_judgment(user_id: str, req: JudgmentRequest):
    if req.decision not in {"approved", "rejected", "pending"}:
        raise HTTPException(status_code=400, detail="Decision harus approved/rejected/pending")
    now = datetime.now(timezone.utc).isoformat()
    update = {
        "judgment_status": req.decision,
        "judgment_note": req.note or "",
        "judged_at": now,
        "judged_by": req.judged_by or "Supervisor Admin",
    }
    result = await db.farmer_profiles.update_one(
        {"user_id": user_id}, {"$set": update}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Petani tidak ditemukan")

    # Auto-invite to next upcoming webinar when approved
    invited_to = None
    if req.decision == "approved":
        upcoming = await db.webinar_events.find_one(
            {"status": "upcoming"}, sort=[("tanggal", 1)]
        )
        if upcoming:
            await db.webinar_events.update_one(
                {"id": upcoming["id"]},
                {"$addToSet": {"invited_user_ids": user_id}},
            )
            await db.farmer_profiles.update_one(
                {"user_id": user_id}, {"$set": {"webinar_invited": True}}
            )
            invited_to = upcoming["id"]

    return {"success": True, "decision": req.decision, "invited_to_webinar": invited_to}


# ----------------------------------------------------
# WEBINAR MANAGEMENT
# ----------------------------------------------------
@api_router.get("/webinars")
async def list_webinars():
    events = await db.webinar_events.find({}, {"_id": 0}).sort("tanggal", 1).to_list(100)
    # enrich with invited farmer names
    for ev in events:
        ids = ev.get("invited_user_ids") or []
        if ids:
            farmers = await db.farmer_profiles.find(
                {"user_id": {"$in": ids}}, {"_id": 0, "user_id": 1, "nama": 1, "desa": 1, "komoditas": 1}
            ).to_list(200)
            ev["invited_farmers"] = farmers
        else:
            ev["invited_farmers"] = []
    return events


@api_router.post("/webinars")
async def create_webinar(payload: WebinarEvent):
    new_id = payload.id or f"wb-{uuid.uuid4().hex[:8]}"
    doc = payload.model_dump()
    doc["id"] = new_id
    doc["created_at"] = datetime.now(timezone.utc).isoformat()
    doc["invited_user_ids"] = doc.get("invited_user_ids") or []
    await db.webinar_events.insert_one(doc.copy())
    return {"success": True, "webinar": doc}


@api_router.delete("/webinars/{webinar_id}")
async def delete_webinar(webinar_id: str):
    r = await db.webinar_events.delete_one({"id": webinar_id})
    return {"success": r.deleted_count > 0}


# ----------------------------------------------------
# BANK PRODUCTS (win-win promo)
# ----------------------------------------------------
@api_router.get("/bank-products")
async def list_bank_products():
    items = await db.bank_products.find({}, {"_id": 0}).sort("plafon_max_rp", -1).to_list(100)
    return items


@api_router.post("/bank-products")
async def create_bank_product(payload: BankProduct):
    new_id = payload.id or f"bp-{uuid.uuid4().hex[:8]}"
    doc = payload.model_dump()
    doc["id"] = new_id
    doc["created_at"] = datetime.now(timezone.utc).isoformat()
    await db.bank_products.insert_one(doc.copy())
    return {"success": True, "product": doc}


@api_router.delete("/bank-products/{product_id}")
async def delete_bank_product(product_id: str):
    r = await db.bank_products.delete_one({"id": product_id})
    return {"success": r.deleted_count > 0}



# ----------------------------------------------------
# SEED DATA ENDPOINT & INITIALIZER
# ----------------------------------------------------
@api_router.post("/seed")
async def seed_database():
    await seed_all_data()
    return {"success": True, "message": "Database ModalTani berhasil diisi data contoh"}

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()    

@app.on_event("startup")
async def startup_event():
    logger.info("Initializing ModalTani seed data on startup...")
    await seed_all_data()


async def seed_all_data():
    # 1. Seed Pre-test Questions
    await db.pretest_questions.delete_many({})
    pretest_data = [
        {
            "id": "pre-1",
            "question": "Mengapa uang hasil panen tidak boleh dicampur langsung dengan uang belanja dapur harian?",
            "category": "Keuangan Rumah Tangga",
            "options": [
                {"id": "a", "text": "Agar kita tahu pasti apakah musim tanam ini untung atau rugi, dan modal tanam berikutnya aman"},
                {"id": "b", "text": "Karena aturan bank mewajibkan punya 5 rekening berbeda"},
                {"id": "c", "text": "Agar uang belanja dapur bisa dihabiskan semuanya"},
                {"id": "d", "text": "Tidak ada pengaruhnya, lebih praktis dicampur satu dompet"}
            ],
            "correct_option_id": "a"
        },
        {
            "id": "pre-2",
            "question": "Berapa suku bunga resmi program Kredit Usaha Rakyat (KUR) Pertanian yang disubsidi pemerintah?",
            "category": "Produk Keuangan Formal",
            "options": [
                {"id": "a", "text": "24% per tahun seperti pinjaman online umum"},
                {"id": "b", "text": "6% efektif per tahun (sekitar 0,2% - 0,5% per bulan)"},
                {"id": "c", "text": "15% per bulan dipotong di awal"},
                {"id": "d", "text": "0% tanpa bunga sama sekali"}
            ],
            "correct_option_id": "b"
        },
        {
            "id": "pre-3",
            "question": "Apa itu skema pembayaran pinjaman 'Yarnen' yang sering digunakan petani tanaman pangan?",
            "category": "Skema Pembayaran",
            "options": [
                {"id": "a", "text": "Bayar angsuran setiap minggu secara tunai ke rentenir"},
                {"id": "b", "text": "Bayar setelah panen tiba (angsuran pokok/bunga dilunasi saat komoditas terjual)"},
                {"id": "c", "text": "Bayar dengan menukar hasil panen gabah langsung ke teller bank"},
                {"id": "d", "text": "Pinjaman yang tidak perlu dibayar jika cuaca buruk"}
            ],
            "correct_option_id": "b"
        },
        {
            "id": "pre-4",
            "question": "Untuk pengajuan KUR Mikro hingga Rp 100 Juta, apakah petani wajib menyerahkan sertifikat tanah sebagai agunan tambahan?",
            "category": "Agunan & Jaminan",
            "options": [
                {"id": "a", "text": "Ya, wajib sertifikat SHM atas nama sendiri"},
                {"id": "b", "text": "Tidak wajib, agunan pokok adalah objek usaha tani dan kelayakan panen"},
                {"id": "c", "text": "Wajib menyerahkan BPKB mobil atau truk"},
                {"id": "d", "text": "Wajib emas batangan 50 gram"}
            ],
            "correct_option_id": "b"
        },
        {
            "id": "pre-5",
            "question": "Dokumen izin usaha apa yang paling mudah diurus petani di tingkat desa untuk syarat administrasi KUR?",
            "category": "Legalitas Administratif",
            "options": [
                {"id": "a", "text": "Surat Izin Usaha Perdagangan (SIUP) PT Besar"},
                {"id": "b", "text": "Surat Keterangan Usaha (SKU) dari Kantor Desa/Kelurahan atau NIB via OSS"},
                {"id": "c", "text": "Akta Notaris Pendirian Perusahaan"},
                {"id": "d", "text": "Izin Ekspor Impor Internasional"}
            ],
            "correct_option_id": "b"
        }
    ]
    await db.pretest_questions.insert_many(pretest_data)

    # 2. Seed PLEK Kementan Learning Modules
    await db.learning_modules.delete_many({})
    modules_data = [
        {
            "id": "mod-1",
            "plek_category": "1. Pengelolaan Keuangan Rumah Tangga",
            "plek_category_number": 1,
            "title": "Pemisahan Dompet Dapur & Dompet Tani",
            "subtitle": "Panduan disiplin kas keluarga petani agar modal tanam berikutnya tidak terpakai kebutuhan konsumtif",
            "duration_minutes": 8,
            "media_type": "video",
            "media_url": "https://www.youtube-nocookie.com/embed/dQw4w9WgXcQ", # Embed standard
            "thumbnail_url": "https://images.unsplash.com/photo-1505471768190-275e2ad7b3f9?w=800&q=80",
            "summary_points": [
                "Pisahkan catatan uang keluarga dan kas operasional usaha tani ke dalam dua amplop/buku berbeda.",
                "Tentukan 'gaji/upah' tetap untuk kebutuhan dapur setiap bulan atau musim, jangan ambil uang modal benih & pupuk.",
                "Sisihkan minimal 10% dari keuntungan bersih panen sebagai dana darurat musim paceklik.",
                "Hindari utang konsumtif barang elektronik dengan bunga harian tinggi."
            ],
            "full_text_notes": "Seringkali petani merasa hasil panen berlimpah tapi uang cepat habis dan bingung saat musim tanam berikutnya tiba karena modal beli pupuk dan bayar buruh tani sudah terpakai belanja dapur. Prinsip PLEK 1 mengajarkan pembagian amplop sederhana: Amplop A (Modal Tanam Berikutnya), Amplop B (Biaya Dapur & Sekolah), Amplop C (Tabungan Darurat Paceklik).",
            "quiz": [
                {
                    "id": "q1-1",
                    "question": "Langkah pertama yang paling praktis untuk memisahkan keuangan keluarga dan usaha tani adalah?",
                    "options": [
                        {"id": "a", "text": "Membuat 2 catatan buku kas/amplop terpisah: Kas Dapur & Kas Tani"},
                        {"id": "b", "text": "Menjual semua hasil panen lalu dibelikan emas langsung"},
                        {"id": "c", "text": "Meminjam uang baru setiap ada kebutuhan mendesak"},
                        {"id": "d", "text": "Membiarkan semua uang masuk ke satu saku celana"}
                    ],
                    "correct_option_id": "a",
                    "explanation": "Pemisahan fisik atau pembukuan sederhana (2 amplop/buku kas) mencegah modal tanam tergerus untuk belanja dapur."
                },
                {
                    "id": "q1-2",
                    "question": "Berapa porsi minimal dari keuntungan bersih panen yang disarankan disimpan sebagai tabungan darurat?",
                    "options": [
                        {"id": "a", "text": "0% karena uang harus selalu diputar"},
                        {"id": "b", "text": "Minimal 10% untuk antisipasi hama/paceklik"},
                        {"id": "c", "text": "100% dan keluarga tidak boleh makan"},
                        {"id": "d", "text": "50% dipinjamkan ke tetangga"}
                    ],
                    "correct_option_id": "b",
                    "explanation": "Dana darurat 10% membantu petani bertahan saat harga anjlok atau ada serangan hama tanpa harus berutang ke rentenir."
                },
                {
                    "id": "q1-3",
                    "question": "Jika ingin membeli perlengkapan rumah tangga baru, dari mana sumber dana yang benar?",
                    "options": [
                        {"id": "a", "text": "Mengambil uang kas modal beli pupuk subsidi"},
                        {"id": "b", "text": "Dari porsi keuntungan bersih keluarga setelah modal tanam aman disisihkan"},
                        {"id": "c", "text": "Mengajukan pinjaman KUR lalu dipakai belanja pribadi"},
                        {"id": "d", "text": "Menjual traktor bantuan kelompok tani"}
                    ],
                    "correct_option_id": "b",
                    "explanation": "KUR dan modal tani hanya boleh digunakan untuk kegiatan produktif pertanian."
                }
            ]
        },
        {
            "id": "mod-2",
            "plek_category": "2. Pengelolaan Keuangan Usaha Tani",
            "plek_category_number": 2,
            "title": "Perhitungan HPP & Arus Kas Musim Tanam",
            "subtitle": "Menghitung modal riil per hektar (benih, pupuk, pestisida, buruh tanam) dan menentukan titik impas harga panen",
            "duration_minutes": 10,
            "media_type": "video",
            "media_url": "https://www.youtube-nocookie.com/embed/dQw4w9WgXcQ",
            "thumbnail_url": "https://images.pexels.com/photos/35544010/pexels-photo-35544010.jpeg?w=800&q=80",
            "summary_points": [
                "Catat 4 komponen utama biaya: Sarana Produksi (Saprotan), Upah Tenaga Kerja, Sewa/Pajak Lahan, dan Biaya Pasca Panen.",
                "Hitung Harga Pokok Produksi (HPP) per kilogram: Total Biaya dibagi Estimasi Hasil Panen (Kg).",
                "Gunakan tabel jadwal pengeluaran mingguan (fase olah tanah, semai, pemupukan 1 & 2, panen).",
                "Pahami titik impas (Break Even Point) agar tidak mudah dipermainkan harga tengkulak nakal."
            ],
            "full_text_notes": "Banyak petani tidak menyadari bahwa upah tenaga kerja keluarga (tenaga sendiri dan istri) dan biaya sewa mesin diesel adalah bagian dari biaya pokok produksi. Dengan mencatat HPP per kg, petani bisa bernegosiasi secara percaya diri saat menjual ke pengepul atau Bulog.",
            "quiz": [
                {
                    "id": "q2-1",
                    "question": "Jika total biaya modal 1 hektar padi adalah Rp 12.000.000 dan hasil panen 3.000 kg gabah, berapa HPP per kg gabah?",
                    "options": [
                        {"id": "a", "text": "Rp 2.000 / kg"},
                        {"id": "b", "text": "Rp 4.000 / kg (Rp 12.000.000 ÷ 3.000 kg)"},
                        {"id": "c", "text": "Rp 6.000 / kg"},
                        {"id": "d", "text": "Rp 10.000 / kg"}
                    ],
                    "correct_option_id": "b",
                    "explanation": "HPP = Total Biaya Operasional dibagi Total Hasil Panen = Rp 12.000.000 / 3.000 kg = Rp 4.000/kg. Jika harga jual di atas Rp 4.000, petani untung!"
                },
                {
                    "id": "q2-2",
                    "question": "Mengapa pencatatan tanggal pengeluaran pupuk dan upah buruh sangat penting saat mengajukan kredit ke bank?",
                    "options": [
                        {"id": "a", "text": "Sebagai bukti kepada analis bank bahwa usaha tani terencana dan memiliki arus kas yang jelas"},
                        {"id": "b", "text": "Hanya untuk formalitas tanpa dicek"},
                        {"id": "c", "text": "Agar mendapat pupuk gratis dari bank"},
                        {"id": "d", "text": "Untuk menaikkan pajak penghasilan"}
                    ],
                    "correct_option_id": "a",
                    "explanation": "Pencatatan yang rapi membuat analis/mantri bank percaya bahwa petani memiliki kemampuan bayar (repayment capacity)."
                }
            ]
        },
        {
            "id": "mod-3",
            "plek_category": "3. Produk & Layanan Keuangan Formal",
            "plek_category_number": 3,
            "title": "Memahami Skema KUR 6% & Prosedur Bank",
            "subtitle": "Kiat lolos wawancara mantri bank, dokumen wajib SKU/NIB, serta pemanfaatan skema pembayaran Yarnen",
            "duration_minutes": 12,
            "media_type": "video",
            "media_url": "https://www.youtube-nocookie.com/embed/dQw4w9WgXcQ",
            "thumbnail_url": "https://images.unsplash.com/photo-1551288049-bebda4e38f71?w=800&q=80",
            "summary_points": [
                "KUR Mikro memberikan pinjaman modal kerja s.d Rp 100 Juta dengan bunga subsidi 6% per tahun.",
                "Pengajuan hingga Rp 100 Juta TANPA agunan tambahan (tidak wajib sertifikat tanah).",
                "Siapkan berkas: KTP, KK, dan Surat Keterangan Usaha (SKU) dari Balai Desa / NIB.",
                "Pilih skema angsuran musiman (Yarnen) jika komoditas Anda adalah tanaman pangan seperti padi atau jagung.",
                "Hindari calo! Pengajuan KUR di bank resmi tidak dipungut biaya pendaftaran atau uang pelicin."
            ],
            "full_text_notes": "Program KUR Pertanian disubsidi langsung oleh Kementerian Keuangan & OJK. Bank penyalur resmi seperti BRI, Mandiri, BNI, dan BSI memiliki mantri lapangan yang bertugas mengecek sawah/kebun secara langsung.",
            "quiz": [
                {
                    "id": "q3-1",
                    "question": "Berapa batas maksimal pinjaman KUR Mikro yang TIDAK memerlukan sertifikat agunan tambahan?",
                    "options": [
                        {"id": "a", "text": "Sampai dengan Rp 100 Juta"},
                        {"id": "b", "text": "Hanya sampai Rp 5 Juta"},
                        {"id": "c", "text": "Maksimal Rp 1 Miliar"},
                        {"id": "d", "text": "Semua pinjaman wajib agunan sertifikat tanah"}
                    ],
                    "correct_option_id": "a",
                    "explanation": "Sesuai Permenko Perekonomian No. 1 Tahun 2023, KUR Mikro hingga plafon Rp 100 Juta tidak diwajibkan agunan tambahan."
                },
                {
                    "id": "q3-2",
                    "question": "Jika seorang calo meminta komisi 10% di muka untuk mencairkan KUR, apa tindakan yang benar?",
                    "options": [
                        {"id": "a", "text": "Memberikannya agar cepat cair"},
                        {"id": "b", "text": "Menolak tegas dan langsung mengajukan sendiri ke kantor bank atau melalui mantri resmi tanpa biaya"},
                        {"id": "c", "text": "Mengajak teman lain untuk ikut calo tersebut"},
                        {"id": "d", "text": "Membayar separuh harga"}
                    ],
                    "correct_option_id": "b",
                    "explanation": "Pengajuan KUR resmi di Bank BUMN GRATIS tanpa biaya calo atau potongan uang pelicin."
                }
            ]
        }
    ]
    await db.learning_modules.insert_many(modules_data)

    # 3. Seed Official KUR Documents Knowledge Base (RAG Source)
    await db.kur_documents.delete_many({})
    kur_kb = [
        {
            "id": "doc-1",
            "topik": "Syarat Umum & Kriteria Calon Debitur KUR",
            "judul": "Pedoman Kriteria Penerima KUR Sektor Pertanian",
            "kategori": "Regulasi Umum",
            "isi_teks": "Calon penerima KUR Pertanian adalah individu/kelompok tani yang memiliki usaha produktif dan layak di sektor tanaman pangan, hortikultura, perkebunan, atau peternakan yang telah berjalan minimal 6 bulan secara aktif. Calon debitur tidak sedang menerima kredit modal kerja komersial lain dari perbankan, namun diperbolehkan memiliki kredit konsumtif seperti KPR atau kredit kendaraan bermotor dengan riwayat pembayaran lancar (Kolektibilitas 1). Usia calon debitur minimal 21 tahun atau sudah menikah.",
            "sumber_nama": "Permenko Perekonomian RI No. 1 Tahun 2023 tentang Pedoman Pelaksanaan KUR",
            "sumber_link": "https://ekon.go.id/publikasi/detail/kur-pedoman",
            "pasal_rujukan": "Pasal 3 & 4 (Kriteria Kelayakan Usaha)"
        },
        {
            "id": "doc-2",
            "topik": "Skema Suku Bunga Subsidi & Pembayaran Yarnen",
            "judul": "Ketentuan Bunga 6% & Skema Angsuran Musiman (Yarnen)",
            "kategori": "Suku Bunga & Angsuran",
            "isi_teks": "Suku bunga KUR Mikro dan KUR Kecil ditetapkan sebesar 6% efektif per tahun untuk pinjaman pertama kali (disubsidi oleh Pemerintah). Untuk sektor pertanian musiman (padi, jagung, bawang, cabai), bank penyalur menyediakan fitur pembayaran fleksibel atau 'Yarnen' (Bayar Setelah Panen) dengan opsi pembayaran bunga secara periodik dan pelunasan pokok saat panen tiba, atau pembayaran pokok dan bunga sekaligus saat musim panen selesai.",
            "sumber_nama": "Otoritas Jasa Keuangan (OJK) & Kementerian Koordinator Bidang Perekonomian",
            "sumber_link": "https://ojk.go.id/id/berita-dan-kegiatan/publikasi/Pages/Panduan-KUR.aspx",
            "pasal_rujukan": "Pasal 12 (Skema Grace Period & Bunga Subsidi)"
        },
        {
            "id": "doc-3",
            "topik": "Jenis Plafon KUR: Super Mikro, Mikro, dan Kecil",
            "judul": "Klasifikasi Batas Plafon Pinjaman KUR Sektor Pertanian",
            "kategori": "Plafon Pinjaman",
            "isi_teks": "1. KUR Super Mikro: Plafon pinjaman sampai dengan Rp 10.000.000 dengan suku bunga 3% per tahun (khusus usaha rintisan/skala kecil).\n2. KUR Mikro: Plafon pinjaman di atas Rp 10.000.000 sampai dengan Rp 100.000.000 dengan suku bunga 6% per tahun.\n3. KUR Kecil: Plafon pinjaman di atas Rp 100.000.000 sampai dengan Rp 500.000.000 dengan suku bunga 6% per tahun (membutuhkan agunan tambahan berupa SHM/BPKB). Jangka waktu modal kerja maksimal 3 tahun atau investasi 5 tahun.",
            "sumber_nama": "Ketentuan Bank Penyalur KUR (Himbara: BRI, Mandiri, BNI, BSI)",
            "sumber_link": "https://bri.co.id/kur",
            "pasal_rujukan": "Petunjuk Teknis Operasional KUR Perbankan"
        },
        {
            "id": "doc-4",
            "topik": "Ketentuan Agunan & Jaminan KUR Tanpa Jaminan",
            "judul": "Regulasi Agunan Pokok & Agunan Tambahan KUR",
            "kategori": "Agunan",
            "isi_teks": "Sesuai Peraturan Menko Perekonomian, untuk pinjaman KUR Super Mikro (s.d Rp 10 juta) dan KUR Mikro (s.d Rp 100 juta), bank dilarang meminta agunan tambahan (seperti sertifikat tanah/rumah atau BPKB kendaraan). Agunan pokok adalah objek yang dibiayai (misalnya tanaman padi yang ditanam, sarana saprotan, atau hasil panen). Penilaian kelayakan kredit difokuskan pada prospek panen, karakter petani, dan kejujuran dalam pembukuan arus kas.",
            "sumber_nama": "Surat Edaran OJK & Permenko Perekonomian No. 1/2023",
            "sumber_link": "https://ojk.go.id",
            "pasal_rujukan": "Pasal 14 Ayat (3) - Larangan Agunan Tambahan KUR Mikro"
        },
        {
            "id": "doc-5",
            "topik": "Dokumen Administratif Wajib",
            "judul": "Daftar Berkas Persyaratan Administrasi Pengajuan KUR",
            "kategori": "Dokumen",
            "isi_teks": "Dokumen yang wajib disiapkan pemohon: 1. Fotokopi e-KTP Pemohon dan Pasangan (jika menikah).\n2. Fotokopi Kartu Keluarga (KK).\n3. Surat Keterangan Usaha (SKU) yang diterbitkan oleh Kepala Desa/Kelurahan setempat ATAU Nomor Induk Berusaha (NIB) berbasis risiko.\n4. NPWP (khusus pinjaman di atas Rp 50 Juta).\n5. Bukti penguasaan lahan (SPPT PBB / Surat Perjanjian Sewa / Surat Keterangan Garap Desa).\n6. Buku rekening tabungan bank penyalur.",
            "sumber_nama": "Standar Operasional Prosedur (SOP) Kredit Mikro Bank BUMN",
            "sumber_link": "https://kemenkopukm.go.id",
            "pasal_rujukan": "Bab III Persyaratan Berkas Calon Debitur"
        },
        {
            "id": "doc-6",
            "topik": "Solusi Status Lahan Sewa / Garapan / Tanah Adat",
            "judul": "Pedoman Pengajuan KUR untuk Petani Penggarap dan Lahan Sewa",
            "kategori": "Legalitas Lahan",
            "isi_teks": "Petani yang tidak memiliki sertifikat tanah pribadi (misalnya petani penggarap, penyewa lahan, atau penggarap tanah adat/kas desa) tetap memenuhi syarat mengajukan KUR Pertanian. Syarat tambahannya adalah melampirkan Surat Perjanjian Sewa/Bagi Hasil yang masih berlaku minimal 1 musim tanam, atau Surat Keterangan Garap yang dilegalisir oleh Kepala Desa/Lurah setempat. Bank akan memverifikasi fisik keberadaan lahan dan tanaman saat survei lapangan.",
            "sumber_nama": "Kementerian Pertanian RI & Direktorat Jenderal Prasarana dan Sarana Pertanian (PSP)",
            "sumber_link": "https://pertanian.go.id",
            "pasal_rujukan": "Juklak Pembiayaan Pertanian Inklusif"
        },
        {
            "id": "doc-7",
            "topik": "Alur & Tahapan Pengajuan KUR di Bank",
            "judul": "Tahapan Lengkap Pengajuan KUR dari Berkas Hingga Pencairan",
            "kategori": "Alur Pengajuan",
            "isi_teks": "Alur pengajuan KUR Pertanian: 1. Registrasi & pengumpulan berkas (KTP, KK, SKU/NIB, bukti lahan).\n2. Pengecekan riwayat SLIK OJK oleh sistem bank.\n3. Kunjungan survei fisik lahan (On-the-Spot) oleh Mantri/Analis Bank untuk mengukur luas lahan dan taksiran hasil panen.\n4. Analisis kelayakan dan persetujuan komite kredit (biasanya 2-5 hari kerja).\n5. Penandatanganan akad kredit dan pembukaan rekening.\n6. Pencairan dana langsung ke rekening petani tanpa potongan calo.",
            "sumber_nama": "Buku Panduan Fasilitasi Pembiayaan KUR Bank Mandiri & BRI",
            "sumber_link": "https://bankmandiri.co.id/kur",
            "pasal_rujukan": "Alur Pelayanan Debitur Mikro"
        },
        {
            "id": "doc-8",
            "topik": "Tips Sukses Lolos Wawancara & Verifikasi Mantri Bank",
            "judul": "Kiat Pra-Skrining dan Wawancara Analis Kredit Pertanian",
            "kategori": "Tips Debitur",
            "isi_teks": "Hal yang dinilai mantri bank saat survei: 1. Karakter kejujuran debitur dan tidak memiliki riwayat kredit macet di tempat lain.\n2. Kejelasan rencana penggunaan dana (hanya untuk sarana produksi dan upah tanam, bukan untuk barang mewah).\n3. Sertifikat atau bukti kelulusan edukasi literasi keuangan (seperti modul ModalTani) menjadi nilai plus mitigasi risiko.\n4. Kesiapan saluran penjualan hasil panen (apakah sudah ada langganan pedagang/koperasi pengumpul).",
            "sumber_nama": "Modul Pendampingan Keuangan Mikro Kementan & ModalTani",
            "sumber_link": "https://modaltani.id",
            "pasal_rujukan": "Modul Mitigasi Risiko Debitur Mikro"
        }
    ]
    await db.kur_documents.insert_many(kur_kb)

    # 4. Seed Demo Users & Farmer Profiles
    await db.users.delete_many({})
    await db.farmer_profiles.delete_many({})
    await db.learning_progress.delete_many({})

    demo_users = [
        {
            "id": "user-budi",
            "nama": "Budi Hartono",
            "role": "petani",
            "no_hp": "0812-8877-6655",
            "desa": "Desa Karanganyar",
            "kecamatan": "Kec. Klari",
            "kabupaten": "Kab. Karawang",
            "komoditas": "Padi Sawah (Ciherang)",
            "avatar_url": "https://images.unsplash.com/photo-1505471768190-275e2ad7b3f9?w=150"
        },
        {
            "id": "user-siti",
            "nama": "Siti Aminah",
            "role": "petani",
            "no_hp": "0857-1122-3344",
            "desa": "Desa Pagersari",
            "kecamatan": "Kec. Candiroto",
            "kabupaten": "Kab. Temanggung",
            "komoditas": "Kopi Robusta & Hortikultura",
            "avatar_url": "https://images.pexels.com/photos/35544010/pexels-photo-35544010.jpeg?w=150"
        },
        {
            "id": "user-joko",
            "nama": "Joko Susilo",
            "role": "petani",
            "no_hp": "0813-9988-7711",
            "desa": "Desa Suntenjaya",
            "kecamatan": "Kec. Lembang",
            "kabupaten": "Kab. Bandung Barat",
            "komoditas": "Hortikultura (Cabai & Tomat)",
            "avatar_url": "https://images.unsplash.com/photo-1534528741775-53994a69daeb?w=150"
        },
        {
            "id": "user-herman",
            "nama": "Herman Wijaya",
            "role": "petani",
            "no_hp": "0821-4455-6677",
            "desa": "Desa Sungai Lilin",
            "kecamatan": "Kec. Sungai Lilin",
            "kabupaten": "Kab. Musi Banyuasin",
            "komoditas": "Karet & Sawit Swadaya",
            "avatar_url": "https://images.unsplash.com/photo-1507003211169-0a1dd7228f2d?w=150"
        },
        {
            "id": "user-dewi",
            "nama": "Dewi Lestari",
            "role": "petani",
            "no_hp": "0878-3344-5566",
            "desa": "Desa Klampok",
            "kecamatan": "Kec. Wanasari",
            "kabupaten": "Kab. Brebes",
            "komoditas": "Bawang Merah",
            "avatar_url": "https://images.unsplash.com/photo-1494790108377-be9c29b29330?w=150"
        },
        {
            "id": "user-wayan",
            "nama": "I Wayan Sudirta",
            "role": "petani",
            "no_hp": "0819-2233-4455",
            "desa": "Desa Jatiluwih",
            "kecamatan": "Kec. Penebel",
            "kabupaten": "Kab. Tabanan",
            "komoditas": "Padi Merah Organik",
            "avatar_url": "https://images.unsplash.com/photo-1500648767791-00dcc994a43e?w=150"
        },
        {
            "id": "user-admin-bri",
            "nama": "Bambang Prakoso (Analis Kredit Mitra Bank)",
            "role": "admin_bank",
            "no_hp": "0811-2233-4400",
            "desa": "Kantor Cabang Mikro",
            "kecamatan": "Regional Hub Jawa Barat",
            "kabupaten": "Bank Mitra Penyalur",
            "komoditas": "Sektor Pembiayaan Pertanian",
            "avatar_url": "https://images.unsplash.com/photo-1472099645785-5658abf4ff4e?w=150"
        }
    ]
    await db.users.insert_many(demo_users)

    demo_profiles = [
        {
            "user_id": "user-budi",
            "nama": "Budi Hartono",
            "nik": "3215081204780002",
            "no_hp": "0812-8877-6655",
            "desa": "Desa Karanganyar",
            "kecamatan": "Kec. Klari",
            "kabupaten": "Kab. Karawang",
            "provinsi": "Jawa Barat",
            "komoditas": "Padi Sawah (Ciherang)",
            "luas_lahan_ha": 2.5,
            "status_lahan": "Milik Sendiri",
            "lama_bertani_tahun": 8,
            "estimasi_pendapatan_musim_rp": 38000000,
            "riwayat_panen": [
                {"musim": "Gadu (Kemarau)", "tahun": 2024, "volume_ton": 13.5, "pendapatan_rp": 36000000},
                {"musim": "Rendeng (Hujan)", "tahun": 2024, "volume_ton": 15.2, "pendapatan_rp": 42000000},
                {"musim": "Gadu (Kemarau)", "tahun": 2025, "volume_ton": 14.0, "pendapatan_rp": 38000000}
            ],
            "dokumen": {
                "ktp": True,
                "kk": True,
                "nib_atau_sku": True,
                "sppt_pbb_atau_surat_lahan": True,
                "buku_tabungan": True,
                "foto_lahan": True
            },
            "pretest_completed": True,
            "pretest_score": 100,
            "recommended_level": "Menengah (Fokus Produk KUR)",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "user_id": "user-siti",
            "nama": "Siti Aminah",
            "nik": "3323055106850001",
            "no_hp": "0857-1122-3344",
            "desa": "Desa Pagersari",
            "kecamatan": "Kec. Candiroto",
            "kabupaten": "Kab. Temanggung",
            "provinsi": "Jawa Tengah",
            "komoditas": "Kopi Robusta & Hortikultura",
            "luas_lahan_ha": 1.2,
            "status_lahan": "Milik Sendiri",
            "lama_bertani_tahun": 5,
            "estimasi_pendapatan_musim_rp": 26000000,
            "riwayat_panen": [
                {"musim": "Panen Raya Kopi", "tahun": 2024, "volume_ton": 3.2, "pendapatan_rp": 28000000},
                {"musim": "Panen Sela Sayur", "tahun": 2024, "volume_ton": 2.1, "pendapatan_rp": 12000000}
            ],
            "dokumen": {
                "ktp": True,
                "kk": True,
                "nib_atau_sku": True,
                "sppt_pbb_atau_surat_lahan": True,
                "buku_tabungan": True,
                "foto_lahan": True
            },
            "pretest_completed": True,
            "pretest_score": 80,
            "recommended_level": "Menengah (Fokus Produk KUR)",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "user_id": "user-joko",
            "nama": "Joko Susilo",
            "nik": "3217031509920004",
            "no_hp": "0813-9988-7711",
            "desa": "Desa Suntenjaya",
            "kecamatan": "Kec. Lembang",
            "kabupaten": "Kab. Bandung Barat",
            "provinsi": "Jawa Barat",
            "komoditas": "Hortikultura (Cabai & Tomat)",
            "luas_lahan_ha": 0.8,
            "status_lahan": "Sewa",
            "lama_bertani_tahun": 3,
            "estimasi_pendapatan_musim_rp": 18000000,
            "riwayat_panen": [
                {"musim": "Musim Cabai 1", "tahun": 2024, "volume_ton": 2.5, "pendapatan_rp": 20000000},
                {"musim": "Musim Tomat 2", "tahun": 2024, "volume_ton": 3.8, "pendapatan_rp": 14000000}
            ],
            "dokumen": {
                "ktp": True,
                "kk": True,
                "nib_atau_sku": False, # Missing SKU
                "sppt_pbb_atau_surat_lahan": True,
                "buku_tabungan": True,
                "foto_lahan": True
            },
            "pretest_completed": True,
            "pretest_score": 60,
            "recommended_level": "Dasar (Kelola Kas & HPP Dulu)",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "user_id": "user-herman",
            "nama": "Herman Wijaya",
            "nik": "1606041201800003",
            "no_hp": "0821-4455-6677",
            "desa": "Desa Sungai Lilin",
            "kecamatan": "Kec. Sungai Lilin",
            "kabupaten": "Kab. Musi Banyuasin",
            "provinsi": "Sumatera Selatan",
            "komoditas": "Karet & Sawit Swadaya",
            "luas_lahan_ha": 3.0,
            "status_lahan": "Milik Sendiri",
            "lama_bertani_tahun": 4,
            "estimasi_pendapatan_musim_rp": 22000000,
            "riwayat_panen": [
                {"musim": "Sadap Karet Semester 1", "tahun": 2024, "volume_ton": 4.5, "pendapatan_rp": 24000000}
            ],
            "dokumen": {
                "ktp": True,
                "kk": True,
                "nib_atau_sku": False,
                "sppt_pbb_atau_surat_lahan": True,
                "buku_tabungan": False,
                "foto_lahan": True
            },
            "pretest_completed": False,
            "pretest_score": 0,
            "recommended_level": "Dasar (Kelola Kas & HPP Dulu)",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "user_id": "user-dewi",
            "nama": "Dewi Lestari",
            "nik": "3329046208950002",
            "no_hp": "0878-3344-5566",
            "desa": "Desa Klampok",
            "kecamatan": "Kec. Wanasari",
            "kabupaten": "Kab. Brebes",
            "provinsi": "Jawa Tengah",
            "komoditas": "Bawang Merah",
            "luas_lahan_ha": 0.4,
            "status_lahan": "Sewa",
            "lama_bertani_tahun": 1,
            "estimasi_pendapatan_musim_rp": 7500000,
            "riwayat_panen": [
                {"musim": "Panen Perdana Bawang", "tahun": 2024, "volume_ton": 1.1, "pendapatan_rp": 7500000}
            ],
            "dokumen": {
                "ktp": True,
                "kk": False,
                "nib_atau_sku": False,
                "sppt_pbb_atau_surat_lahan": False,
                "buku_tabungan": False,
                "foto_lahan": True
            },
            "pretest_completed": False,
            "pretest_score": 0,
            "recommended_level": "Dasar (Kelola Kas & HPP Dulu)",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "user_id": "user-wayan",
            "nama": "I Wayan Sudirta",
            "nik": "5102041903820001",
            "no_hp": "0819-2233-4455",
            "desa": "Desa Jatiluwih",
            "kecamatan": "Kec. Penebel",
            "kabupaten": "Kab. Tabanan",
            "provinsi": "Bali",
            "komoditas": "Padi Merah Organik",
            "luas_lahan_ha": 1.6,
            "status_lahan": "Milik Sendiri",
            "lama_bertani_tahun": 10,
            "estimasi_pendapatan_musim_rp": 45000000,
            "riwayat_panen": [
                {"musim": "Panen Padi Organik 1", "tahun": 2024, "volume_ton": 8.0, "pendapatan_rp": 48000000},
                {"musim": "Panen Padi Organik 2", "tahun": 2024, "volume_ton": 7.5, "pendapatan_rp": 44000000}
            ],
            "dokumen": {
                "ktp": True,
                "kk": True,
                "nib_atau_sku": True,
                "sppt_pbb_atau_surat_lahan": True,
                "buku_tabungan": True,
                "foto_lahan": True
            },
            "pretest_completed": True,
            "pretest_score": 100,
            "recommended_level": "Menengah (Fokus Produk KUR)",
            "created_at": datetime.now(timezone.utc).isoformat()
        }
    ]
    await db.farmer_profiles.insert_many(demo_profiles)

    # 5. Seed Initial Learning Progress for Budi & Wayan & Siti
    demo_progress = [
        # Pak Budi completed all 3 modules
        {"user_id": "user-budi", "module_id": "mod-1", "status": "completed", "quiz_score": 100, "updated_at": datetime.now(timezone.utc).isoformat()},
        {"user_id": "user-budi", "module_id": "mod-2", "status": "completed", "quiz_score": 100, "updated_at": datetime.now(timezone.utc).isoformat()},
        {"user_id": "user-budi", "module_id": "mod-3", "status": "completed", "quiz_score": 100, "updated_at": datetime.now(timezone.utc).isoformat()},
        # Ibu Siti completed 2 modules
        {"user_id": "user-siti", "module_id": "mod-1", "status": "completed", "quiz_score": 100, "updated_at": datetime.now(timezone.utc).isoformat()},
        {"user_id": "user-siti", "module_id": "mod-2", "status": "completed", "quiz_score": 80, "updated_at": datetime.now(timezone.utc).isoformat()},
        # Pak Joko completed 1 module
        {"user_id": "user-joko", "module_id": "mod-1", "status": "completed", "quiz_score": 75, "updated_at": datetime.now(timezone.utc).isoformat()},
        # Pak Wayan completed all 3 modules
        {"user_id": "user-wayan", "module_id": "mod-1", "status": "completed", "quiz_score": 100, "updated_at": datetime.now(timezone.utc).isoformat()},
        {"user_id": "user-wayan", "module_id": "mod-2", "status": "completed", "quiz_score": 100, "updated_at": datetime.now(timezone.utc).isoformat()},
        {"user_id": "user-wayan", "module_id": "mod-3", "status": "completed", "quiz_score": 100, "updated_at": datetime.now(timezone.utc).isoformat()}
    ]
    await db.learning_progress.insert_many(demo_progress)
    logger.info("ModalTani seeding completed successfully.")

    # Seed Webinar Events
    await db.webinar_events.delete_many({})
    now_iso = datetime.now(timezone.utc).isoformat()
    webinars = [
        {
            "id": "wb-pemantapan-01",
            "judul": "Pemantapan Pengajuan KUR Pertanian Musim Tanam 2026",
            "deskripsi": "Sesi tanya-jawab langsung dengan analis Bank Himbara & narasumber Kementerian. Peserta akan dipandu simulasi berkas hingga siap datang ke unit bank.",
            "tanggal": "2026-03-15T09:00:00+07:00",
            "durasi_menit": 120,
            "lokasi": "Zoom Meeting (Link dikirim ke WhatsApp peserta)",
            "link_daftar": "https://modaltani.id/webinar/pemantapan-kur-mar2026",
            "pembicara": [
                {"nama": "Ir. Retno Wulandari", "jabatan": "Kepala Sub-Direktorat Pembiayaan Petani", "institusi": "Kementerian Pertanian RI"},
                {"nama": "Dr. Bambang Susanto", "jabatan": "Deputi Direktur Kredit Program", "institusi": "OJK"},
                {"nama": "Rina Kartika, S.E., M.M.", "jabatan": "Regional Head Mikro Banking", "institusi": "Bank BRI"},
                {"nama": "Ahmad Fauzi", "jabatan": "Kepala Unit KUR Pertanian", "institusi": "Bank Mandiri"},
            ],
            "invited_user_ids": ["user-budi", "user-wayan"],
            "status": "upcoming",
            "created_at": now_iso,
        },
        {
            "id": "wb-konsultasi-1on1",
            "judul": "Konsultasi 1-on-1 KUR Kecil (> Rp 100 Juta)",
            "deskripsi": "Sesi eksklusif untuk calon debitur KUR Kecil dengan Kepala Cabang. Wajib membawa BPJS Ketenagakerjaan dan bukti agunan tambahan.",
            "tanggal": "2026-04-05T13:00:00+07:00",
            "durasi_menit": 60,
            "lokasi": "Hybrid (Kantor Cabang / Zoom Room)",
            "link_daftar": "https://modaltani.id/webinar/kur-kecil-apr2026",
            "pembicara": [
                {"nama": "Ir. Djoko Prasetyo", "jabatan": "Kepala Cabang Utama", "institusi": "Bank BNI"},
                {"nama": "Sri Handayani, S.E.", "jabatan": "Direktur Kredit Retail", "institusi": "Bank BSI"},
            ],
            "invited_user_ids": [],
            "status": "upcoming",
            "created_at": now_iso,
        },
    ]
    await db.webinar_events.insert_many(webinars)

    # Seed Bank Products
    await db.bank_products.delete_many({})
    bank_products = [
        {
            "id": "bp-bri-mikro",
            "bank_nama": "Bank BRI",
            "bank_logo_url": "https://upload.wikimedia.org/wikipedia/commons/thumb/2/2e/BANK_BRI_logo.svg/512px-BANK_BRI_logo.svg.png",
            "nama_produk": "KUR Mikro BRI Pertanian",
            "tier": "KUR Mikro",
            "bunga_efektif_tahun": 6.0,
            "plafon_min_rp": 10_000_000,
            "plafon_max_rp": 100_000_000,
            "tenor_bulan_max": 36,
            "highlight": "Tanpa agunan tambahan, mantri datang ke sawah, skema Yarnen tersedia",
            "link": "https://bri.co.id/kur",
            "kontak_mantri": "1500-017",
            "active": True,
            "created_at": now_iso,
        },
        {
            "id": "bp-mandiri-super",
            "bank_nama": "Bank Mandiri",
            "bank_logo_url": "https://upload.wikimedia.org/wikipedia/commons/thumb/a/ad/Bank_Mandiri_logo_2016.svg/512px-Bank_Mandiri_logo_2016.svg.png",
            "nama_produk": "KUR Super Mikro Mandiri Tani",
            "tier": "KUR Super Mikro",
            "bunga_efektif_tahun": 3.0,
            "plafon_min_rp": 1_000_000,
            "plafon_max_rp": 10_000_000,
            "tenor_bulan_max": 24,
            "highlight": "Cocok untuk petani pemula, cair 3 hari kerja",
            "link": "https://bankmandiri.co.id/kur",
            "kontak_mantri": "14000",
            "active": True,
            "created_at": now_iso,
        },
        {
            "id": "bp-bni-kecil",
            "bank_nama": "Bank BNI",
            "bank_logo_url": "https://upload.wikimedia.org/wikipedia/commons/thumb/5/55/BNI_logo.svg/512px-BNI_logo.svg.png",
            "nama_produk": "KUR Kecil BNI Agri Prima",
            "tier": "KUR Kecil",
            "bunga_efektif_tahun": 6.0,
            "plafon_min_rp": 101_000_000,
            "plafon_max_rp": 500_000_000,
            "tenor_bulan_max": 60,
            "highlight": "Untuk agrobisnis berkembang. Wajib BPJS TK & agunan SHM/BPKB. Konsultasi 1-on-1 dengan Kepala Cabang",
            "link": "https://bni.co.id/id-id/beranda/produk/mikro/kur-kecil",
            "kontak_mantri": "1500046",
            "active": True,
            "created_at": now_iso,
        },
        {
            "id": "bp-bsi-syariah",
            "bank_nama": "Bank BSI",
            "bank_logo_url": "https://upload.wikimedia.org/wikipedia/id/thumb/e/e4/Bank_Syariah_Indonesia.svg/512px-Bank_Syariah_Indonesia.svg.png",
            "nama_produk": "KUR Mikro iB Hasanah",
            "tier": "KUR Mikro",
            "bunga_efektif_tahun": 6.0,
            "plafon_min_rp": 5_000_000,
            "plafon_max_rp": 100_000_000,
            "tenor_bulan_max": 36,
            "highlight": "Akad Murabahah (jual-beli) — sesuai prinsip syariah, cocok untuk petani muslim",
            "link": "https://bankbsi.co.id/produk-layanan/produk/pembiayaan-kur",
            "kontak_mantri": "14040",
            "active": True,
            "created_at": now_iso,
        },
    ]
    await db.bank_products.insert_many(bank_products)

    # Sync knowledge base into Qdrant vector store
    if HAS_RAG:
        try:
            rag_clear()
            n = rag_upsert(kur_kb)
            logger.info(f"[RAG] seeded Qdrant with {n} points")
        except Exception as e:
            logger.warning(f"[RAG] Qdrant seed skipped: {e}")